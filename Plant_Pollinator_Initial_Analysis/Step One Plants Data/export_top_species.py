"""Export top species rankings to Parquet and Word formats."""

from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches

LOG_PATH = Path("species_tally_top55.log")
PARQUET_PATH = Path("top55_species.parquet")
DOCX_PATH = Path("top55_species.docx")


def parse_log(log_path: Path) -> tuple[pd.DataFrame, dict[str, str | int | float]]:
    text = log_path.read_text()
    summary: dict[str, str | int | float] = {}

    finished = re.search(
        r"Finished\. Total rows: ([\d,]+) \| Unique species: ([\d,]+)",
        text,
    )
    if finished:
        summary["total_rows"] = int(finished.group(1).replace(",", ""))
        summary["unique_species"] = int(finished.group(2).replace(",", ""))

    cumulative = re.search(
        r"Cumulative records in top 55 species: ([\d,]+) \(([\d.]+)% of ([\d,]+) total rows\)",
        text,
    )
    if cumulative:
        summary["cumulative_top55_records"] = int(cumulative.group(1).replace(",", ""))
        summary["cumulative_top55_pct"] = float(cumulative.group(2))
        summary["cumulative_top55_total_rows"] = int(cumulative.group(3).replace(",", ""))

    rank55 = re.search(
        r"55th most populated species: (.+) \(([\d,]+) records\)",
        text,
    )
    if rank55:
        summary["rank_55_species"] = rank55.group(1).strip()
        summary["rank_55_count"] = int(rank55.group(2).replace(",", ""))

    rows: list[dict[str, object]] = []
    for match in re.finditer(
        r"^\s*(\d+)\.\s+(.+?)\s+Count:\s+([\d,]+)\s*$",
        text,
        re.MULTILINE,
    ):
        rows.append(
            {
                "rank": int(match.group(1)),
                "species": match.group(2).strip(),
                "count": int(match.group(3).replace(",", "")),
            }
        )

    if not rows:
        raise ValueError(f"No ranked species found in {log_path}")

    df = pd.DataFrame(rows).sort_values("rank").reset_index(drop=True)
    return df, summary


def export_parquet(df: pd.DataFrame, summary: dict[str, str | int | float], path: Path) -> None:
    metadata = {k: str(v) for k, v in summary.items()}
    df.to_parquet(path, index=False, compression="snappy")
    print(f"Wrote {path} ({len(df)} rows, metadata: {metadata})")


def export_docx(df: pd.DataFrame, summary: dict[str, str | int | float], path: Path) -> None:
    doc = Document()
    doc.add_heading("Top 55 Most Populated Species — PhenoField Train Set", level=1)
    doc.add_paragraph(
        "Dataset: dcher95/phenofield (train split). "
        "Species frequencies tallied by streaming 107 shards."
    )

    if summary:
        doc.add_heading("Summary", level=2)
        if "total_rows" in summary:
            doc.add_paragraph(f"Total rows processed: {summary['total_rows']:,}")
        if "unique_species" in summary:
            doc.add_paragraph(f"Unique species: {summary['unique_species']:,}")
        if "cumulative_top55_records" in summary:
            doc.add_paragraph(
                f"Cumulative records in top 55 species: "
                f"{summary['cumulative_top55_records']:,} "
                f"({summary.get('cumulative_top55_pct', 0):.1f}% of dataset)"
            )
        if "rank_55_species" in summary:
            doc.add_paragraph(
                f"55th most populated species: {summary['rank_55_species']} "
                f"({summary.get('rank_55_count', 0):,} records)"
            )

    doc.add_heading("Ranked Species List", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Rank"
    headers[1].text = "Species"
    headers[2].text = "Count"

    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(int(row["rank"]))
        cells[1].text = str(row["species"])
        cells[2].text = f"{int(row['count']):,}"

    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(1.5)

    doc.save(path)
    print(f"Wrote {path}")


def main() -> None:
    if not LOG_PATH.exists():
        raise FileNotFoundError(f"Log file not found: {LOG_PATH}")

    df, summary = parse_log(LOG_PATH)
    export_parquet(df, summary, PARQUET_PATH)
    export_docx(df, summary, DOCX_PATH)


if __name__ == "__main__":
    main()
